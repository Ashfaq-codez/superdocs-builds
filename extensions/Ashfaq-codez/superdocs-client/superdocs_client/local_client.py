import uuid
import datetime
from pathlib import Path
from typing import List, Dict, Any

from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from superdocs_client.interface import SuperDocsClientInterface
from superdocs_client.models import (
    DocumentRef, EditProposal, ApprovalDecision, ApprovalResult, 
    ExportFormat, ExportResult, ExportArtifact
)
from superdocs_client.exceptions import SuperDocsError, ExportError


class LocalSuperDocsClient(SuperDocsClientInterface):
    """Concrete implementation of the SDK that generates real local files."""

    def __init__(self, runtime_dir: Path):
        self.runtime_dir = Path(runtime_dir)
        self.runtime_dir.mkdir(parents=True, exist_ok=True)
        
        self.documents: Dict[str, Dict[str, Any]] = {}
        self.proposals: Dict[str, Dict[str, Any]] = {}

    async def upload(self, file_path: str) -> DocumentRef:
        doc_id = uuid.uuid4().hex[:12]
        self.documents[doc_id] = {"original_file": file_path}
        return DocumentRef(document_id=doc_id)

    async def propose_edit(self, document: DocumentRef, instruction: str) -> EditProposal:
        prop_id = uuid.uuid4().hex[:12]
        self.proposals[prop_id] = {
            "document_id": document.document_id,
            "instruction": instruction,
            "status": "PENDING"
        }
        return EditProposal(
            proposal_id=prop_id, 
            document_id=document.document_id, 
            proposed_changes={"text": instruction}, 
            status="PENDING"
        )

    async def approve_edit(self, proposal_id: str, decision: ApprovalDecision) -> ApprovalResult:
        if proposal_id not in self.proposals:
            raise SuperDocsError(f"Proposal {proposal_id} not found.")
            
        status = "APPROVED" if decision == ApprovalDecision.APPROVE else "REJECTED"
        self.proposals[proposal_id]["status"] = status
        
        return ApprovalResult(proposal_id=proposal_id, status=status, applied=(status == "APPROVED"))

    async def export(self, document: DocumentRef, formats: List[ExportFormat]) -> ExportResult:
        doc_id = document.document_id
        
        approved_text = None
        for p in self.proposals.values():
            if p["document_id"] == doc_id and p["status"] == "APPROVED":
                approved_text = p["instruction"]
                break
                
        if not approved_text:
            raise ExportError(f"Cannot export document {doc_id}: No approved edits found.")

        export_dir = self.runtime_dir / doc_id
        export_dir.mkdir(parents=True, exist_ok=True)
        artifacts = []

        current_date = datetime.date.today().strftime('%B %d, %Y')

        # 1. Generate Professional DOCX
        if ExportFormat.DOCX in formats:
            docx_path = export_dir / "offer_letter.docx"
            doc = Document()
            
            doc.add_heading('SuperDocs Inc.', 0)
            doc.add_paragraph(f"Date: {current_date}")
            
            for line in approved_text.split('\n'):
                clean_line = line.strip()
                if not clean_line:
                    continue
                if clean_line.startswith('# '):
                    doc.add_heading(clean_line[2:], level=1)
                elif clean_line.startswith('## '):
                    doc.add_heading(clean_line[3:], level=2)
                else:
                    doc.add_paragraph(clean_line)
                    
            doc.save(str(docx_path))
            artifacts.append(ExportArtifact(format=ExportFormat.DOCX, reference=f"{doc_id}/offer_letter.docx"))

        # 2. Generate Professional PDF
        if ExportFormat.PDF in formats:
            pdf_path = export_dir / "offer_letter.pdf"
            pdf = FPDF()
            pdf.add_page()
            
            # Header
            pdf.set_font("Helvetica", style="B", size=20)
            pdf.cell(w=0, h=10, text="SuperDocs Inc.", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="L")
            pdf.set_font("Helvetica", size=10)
            pdf.cell(w=0, h=10, text=f"Date: {current_date}", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="L")
            
            # Replaced pdf.ln() with an explicit spacer cell to protect cursor state
            pdf.cell(w=0, h=5, text="", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            safe_text = approved_text.encode('latin-1', 'replace').decode('latin-1')
            for line in safe_text.split('\n'):
                clean_line = line.strip()
                if not clean_line:
                    pdf.cell(w=0, h=4, text="", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    continue
                
                # Dynamic markdown parsing
                if clean_line.startswith('# '):
                    pdf.set_font("Helvetica", style="B", size=16)
                    pdf.cell(w=0, h=8, text=clean_line[2:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.cell(w=0, h=2, text="", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                elif clean_line.startswith('## '):
                    pdf.set_font("Helvetica", style="B", size=12)
                    pdf.cell(w=0, h=6, text=clean_line[3:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.cell(w=0, h=2, text="", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                else:
                    pdf.set_font("Helvetica", size=10)
                    pdf.multi_cell(w=0, h=5, text=clean_line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    
            pdf.output(str(pdf_path))
            artifacts.append(ExportArtifact(format=ExportFormat.PDF, reference=f"{doc_id}/offer_letter.pdf"))

        return ExportResult(document_id=doc_id, artifacts=artifacts, status="EXPORTED")